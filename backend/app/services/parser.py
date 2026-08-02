import re
import os
import fitz  # PyMuPDF
import pdfplumber
import docx
from typing import Dict, Any, List

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    # Try PyMuPDF first
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"PyMuPDF failed: {e}. Trying pdfplumber...")
        # Fallback to pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as ex:
            print(f"pdfplumber failed: {ex}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"python-docx failed: {e}")
    return text

def parse_resume_data(text: str) -> Dict[str, Any]:
    """
    Strips and structures the text using standard regexes and section matchers.
    Can be enhanced using OpenAI API in the AI service.
    """
    result = {
        "name": "",
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
        "portfolio": "",
        "skills": [],
        "education": [],
        "experience": [],
        "projects": [],
        "certifications": [],
        "languages": [],
        "achievements": []
    }

    if not text:
        return result

    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # 1. Simple Name Extractor (often the first line)
    if lines:
        for line in lines[:3]:
            # Simple rule: name is typically 2-3 alphabetic words and not a link or contact info
            if re.match(r"^[A-Z][a-zA-Z]+(\s[A-Z][a-zA-Z]+){1,2}$", line):
                result["name"] = line
                break
        if not result["name"] and lines:
            result["name"] = lines[0]

    # 2. Contact details
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    if email_match:
        result["email"] = email_match.group(0)

    phone_match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    if phone_match:
        result["phone"] = phone_match.group(0)

    # 3. Links
    linkedin_match = re.search(r"(https?://)?(www\.)?linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    if linkedin_match:
        result["linkedin"] = linkedin_match.group(0)

    github_match = re.search(r"(https?://)?(www\.)?github\.com/[\w\-]+", text, re.IGNORECASE)
    if github_match:
        result["github"] = github_match.group(0)

    portfolio_match = re.search(r"(https?://)?(www\.)?([a-zA-Z0-9\-]+\.)+(com|org|net|io|me|dev)/?([\w\-]+)?", text, re.IGNORECASE)
    if portfolio_match:
        # Prevent picking up email domains or github/linkedin
        port = portfolio_match.group(0)
        if "linkedin.com" not in port.lower() and "github.com" not in port.lower():
            result["portfolio"] = port

    # 4. Skills extraction by scanning a tech list
    skill_keywords = [
        "Python", "JavaScript", "TypeScript", "React", "Angular", "Vue", "Node.js", "Express", 
        "Django", "Flask", "FastAPI", "SQL", "PostgreSQL", "MongoDB", "MySQL", "SQLite", 
        "HTML", "CSS", "Tailwind CSS", "Bootstrap", "Git", "GitHub", "Docker", "Kubernetes", 
        "AWS", "Google Cloud", "Azure", "Linux", "C++", "Java", "C#", "Go", "Rust", "Swift", 
        "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "NLP", "spaCy", "Scikit-Learn",
        "Data Analysis", "Redux", "GraphQL", "REST API", "CI/CD", "Terraform", "Agile", "Scrum",
        "Next.js", "Vite", "Solidity", "Framer Motion", "OpenAI", "LLMs", "LangChain"
    ]
    for skill in skill_keywords:
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text, re.IGNORECASE):
            result["skills"].append(skill)

    # 5. Section-based Extraction (Education, Experience, Projects)
    sections = {
        "education": ["education", "academic background", "qualification", "studies"],
        "experience": ["experience", "employment history", "work history", "professional experience"],
        "projects": ["projects", "personal projects", "key projects", "academic projects"],
        "certifications": ["certifications", "certs", "licenses", "courses"],
        "languages": ["languages", "language proficiency"],
        "achievements": ["achievements", "accomplishments", "awards", "honors"]
    }

    current_section = None
    section_buffer = []

    for line in lines:
        lower_line = line.lower()
        matched_section = None
        for sec_name, keywords in sections.items():
            for keyword in keywords:
                if re.match(r"^(" + keyword + r")$", lower_line) or (len(lower_line) < 30 and lower_line.startswith(keyword) and ":" not in lower_line):
                    matched_section = sec_name
                    break
            if matched_section:
                break

        if matched_section:
            if current_section and section_buffer:
                result[current_section] = section_buffer
            current_section = matched_section
            section_buffer = []
        elif current_section:
            section_buffer.append(line)

    if current_section and section_buffer:
        result[current_section] = section_buffer

    # Format output arrays if they are lists of strings
    # For education, experience, and projects we can format them as lists of dictionaries or strings
    # Let's keep them as clean lists of lines or groups of lines.
    return result
